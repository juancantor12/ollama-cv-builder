"""Dox Factory unit testing."""

from pathlib import Path
from unittest.mock import patch, MagicMock
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from resume_generator.docx_factory import DocxFactory


@patch("docx.document.Document.add_paragraph")
def test_add_text(mock_add_paragraph):
    """
    Tests that add_text returns a paragraph with the specified contents.
    This test mocks add_paragraph but not add_run, to be able to retrieve the run attrs
    """
    docx_factory = DocxFactory()
    magick_mock_paragraph = MagicMock()
    mock_add_paragraph.return_value = magick_mock_paragraph
    paragraph = docx_factory.add_text(
        text="Hello_world!",
        centered=True,
        uppercase=True,
        bold=False,
        italic=False,
        font_size=15,
    )
    mock_add_paragraph.assert_called_once()
    # The next assert implicitly tests that paragraph is an instance of docx.text.paragraph.Paragraph
    assert paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
    magick_mock_paragraph.add_run.assert_called_once_with("HELLO WORLD!")
    mock_run = magick_mock_paragraph.add_run.return_value
    assert mock_run.bold is False
    assert mock_run.italic is False
    assert mock_run.font.size == Pt(15)


def test_bulleted():
    """Tests that a given string is returned in bullet format, just for test coverage."""
    assert "  •  bullet" == DocxFactory.bulleted("bullet")


@patch("docx.document.Document.add_paragraph")
def test_add_line_breaks(mock_add_paragraph):
    """Tests that add_line_breaks adds empty paragraphs x times."""
    docx_factory = DocxFactory()
    docx_factory.add_line_breaks(3)
    assert mock_add_paragraph.call_count == 3


@patch("docx.document.Document.save")
def test_save_document(mock_save):
    """Tests that add_line_breaks adds empty paragraphs x times."""
    docx_factory = DocxFactory()
    test_output_path = Path("output/example.com")
    file_path = docx_factory.save_document(test_output_path)
    expected_file_path = test_output_path / "generated_resume.docx"
    mock_save.assert_called_once_with(expected_file_path)
    assert file_path == expected_file_path
