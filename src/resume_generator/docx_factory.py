"""Utilities for the generation of the docx file."""
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

class DocxFactory:
	"""Utilities for the generation of the docx file."""
	def __init__(self):
		self.doc = Document()
		self.font_size = 10.5
		self.style = self.doc.styles['Normal']
		self.style.font.name = "Calibri"
		self.style.font.size = Pt(self.font_size)
		self.style.paragraph_format.line_spacing = 1.0
		self.style.paragraph_format.space_after = Pt(2)
		self.margin = 0.7
		self.section = self.doc.sections[0]
		self.section.top_margin = Inches(self.margin)
		self.section.bottom_margin = Inches(self.margin)
		self.section.left_margin = Inches(self.margin)
		self.section.right_margin = Inches(self.margin)

	def add_text(self, **kwargs) -> Paragraph:
		"""
		Adds a text paragraph: optional parameters for centered, upper_case, bold, italic and font size
		Replaces _ with blank spaces.
		"""
		paragraph = self.doc.add_paragraph()
		if 'centered' in kwargs and kwargs['centered']:
			paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
		text = kwargs['text'].replace("_", " ")
		if kwargs['uppercase']:
			text = text.upper()
		run = paragraph.add_run(text)
		run.bold = kwargs['bold'] if 'bold' in kwargs else False
		run.italic = kwargs['italic'] if 'italic' in kwargs else False
		run.font.size = Pt(kwargs['font_size']) if 'font_size' in kwargs else self.font_size
		return paragraph

	@staticmethod
	def bulleted(bullet: str) -> str:
		"""Adds bullet appearance to a text."""
		return f"  •  {bullet}"

	def add_horizontal_line(self) -> None:
		"""
		Adds an horizontal, black line. 
		python-docx doesn't provide a way to add lines so adding them through the XML layer 
		is the only way, thus, the pylint disable comment
		"""
		paragraph = self.doc.add_paragraph("")
		run = paragraph.add_run()
		run.font.size = Pt(1)
		p_paragraph = paragraph._p 					# pylint: disable=protected-access
		p_borders = OxmlElement('w:pBdr')
		bottom_border = OxmlElement('w:bottom')
		bottom_border.set(qn('w:val'), 'single')
		bottom_border.set(qn('w:sz'), '6')	   		# Thickness
		bottom_border.set(qn('w:space'), '1')		# Space between text and line
		bottom_border.set(qn('w:color'), '000000')  # Black line
		p_borders.append(bottom_border)
		p_pr = p_paragraph.get_or_add_pPr()
		p_pr.append(p_borders)
		self.doc.add_paragraph("")

	def add_line_breaks(self, amount: int) -> None:
		"""Adds an empty paragraph to force a line break."""
		for _ in range(amount):
			self.doc.add_paragraph()

	def save_document(self, output_path: str) -> str:
		"""Saves the docx to the specified output path and returns the file path."""
		file_path = output_path + "/resume.docx"
		self.doc.save(file_path)
		return file_path
